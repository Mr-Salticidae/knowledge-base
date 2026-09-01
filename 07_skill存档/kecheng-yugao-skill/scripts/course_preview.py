# -*- coding: utf-8 -*-
"""Deterministic helpers for the kecheng-yugao-skill."""

from __future__ import annotations

import argparse
import hashlib
import json
import re
import sys
import unicodedata
from datetime import date, datetime, timedelta
from difflib import SequenceMatcher
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor


SCHEMA_VERSION = 2
LIVE = "live"
RECORDED = "recorded"
DELIVERIES = {LIVE, RECORDED}
DELIVERY_EVIDENCE_SOURCES = {
    "schedule_legend",
    "explicit_schedule_label",
    "user_confirmation",
}
TEMPLATE_ID_RE = re.compile(r"^[a-z0-9]+(?:-[a-z0-9]+)*$")
BRACKET_RE = re.compile(r"【\s*(.+?)\s*】")
TEACHER_RE = re.compile(r"(?:授课老师|讲师)\s*[：:]\s*(.+)")
WEEKLY_RE = re.compile(
    r"^(?P<sm>\d{1,2})\.(?P<sd>\d{1,2})-(?P<em>\d{1,2})\.(?P<ed>\d{1,2})课程预告(?:_.+)?\.docx$",
    re.I,
)
CANONICAL_WEEKLY_RE = re.compile(
    r"^\d{1,2}\.\d{1,2}-\d{1,2}\.\d{1,2}课程预告\.docx$",
    re.I,
)


def parse_date(value: str) -> date:
    for fmt in ("%Y-%m-%d", "%Y/%m/%d", "%m.%d", "%m/%d"):
        try:
            parsed = datetime.strptime(value, fmt).date()
            if "%Y" not in fmt:
                parsed = parsed.replace(year=date.today().year)
            return parsed
        except ValueError:
            pass
    raise ValueError(f"Unsupported date: {value!r}")


def period_for(target: date) -> dict[str, Any]:
    if target.weekday() >= 5:
        raise ValueError("Weekend target requires explicit schedule evidence; edit the manifest manually.")
    friday = target + timedelta(days=4 - target.weekday())
    filename = f"{target.month}.{target.day}-{friday.month}.{friday.day}课程预告.docx"
    return {
        "start": target.isoformat(),
        "end": friday.isoformat(),
        "filename": filename,
        "existing_document": None,
    }


def existing_weekly_for(docs: Path, target: date) -> dict[str, Any] | None:
    candidates = []
    for path in docs.glob("*.docx"):
        match = WEEKLY_RE.match(path.name)
        if not match:
            continue
        values = {key: int(value) for key, value in match.groupdict().items()}
        start = date(target.year, values["sm"], values["sd"])
        end_year = target.year + (1 if values["em"] < values["sm"] else 0)
        end = date(end_year, values["em"], values["ed"])
        if start <= target <= end:
            candidates.append((start, end, path))
    if not candidates:
        return None
    if len(candidates) > 1:
        raise ValueError(
            "multiple course-preview DOCX files cover the target date; "
            "resolve the authoritative weekly document first: "
            + " | ".join(item[2].name for item in candidates)
        )
    start, end, path = candidates[0]
    return {
        "start": start.isoformat(),
        "end": end.isoformat(),
        "filename": f"{start.month}.{start.day}-{end.month}.{end.day}课程预告.docx",
        "existing_document": str(path.resolve()),
    }


def normalize_title(value: str) -> str:
    text = unicodedata.normalize("NFKC", value).lower()
    text = re.sub(r"\b(?:ai|al)\b", "ai", text)
    text = text.translate(str.maketrans({"“": "", "”": "", "‘": "", "’": ""}))
    return re.sub(r"[\W_]+", "", text, flags=re.UNICODE)


def course_key_for(course: dict[str, Any]) -> str:
    """Return a stable identity that cannot drift with DOCX paragraph order."""
    class_key = normalize_title(str(course.get("class_name", "")))
    title_key = normalize_title(
        str(course.get("schedule_title") or course.get("course_title") or "")
    )
    return f"{class_key}::{title_key}"


def poster_filename_contract_errors(
    course: dict[str, Any], poster_path: Path | str, target_date: str
) -> list[str]:
    """Validate that a poster filename semantically identifies its course."""
    path = Path(poster_path)
    stem = path.stem
    stem_key = normalize_title(stem)
    errors: list[str] = []
    for field in ("class_name", "course_title", "teacher"):
        value = str(course.get(field) or "")
        if not value or normalize_title(value) not in stem_key:
            errors.append(f"filename is missing {field}={value!r}")

    parsed_date = date.fromisoformat(str(target_date))
    normalized_stem = unicodedata.normalize("NFKC", stem).lower()
    date_tokens = {
        f"{parsed_date.month}.{parsed_date.day}",
        f"{parsed_date.month:02d}.{parsed_date.day:02d}",
        f"{parsed_date.month}.{parsed_date.day:02d}",
        f"{parsed_date.month:02d}.{parsed_date.day}",
    }
    if not any(token in normalized_stem for token in date_tokens):
        errors.append(
            "filename is missing target date token "
            f"{parsed_date.month}.{parsed_date.day}"
        )
    if course.get("delivery") == LIVE and "直播" not in stem:
        errors.append("live poster filename is missing 直播 marker")
    return errors


def read_json(path: Path) -> dict[str, Any]:
    # Windows PowerShell 5 may emit a UTF-8 BOM even when a JSON-producing
    # command requested UTF-8. Accept it at workflow boundaries while continuing
    # to write canonical BOM-free UTF-8 ourselves.
    return json.loads(path.read_text(encoding="utf-8-sig"))


def write_json(path: Path, data: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")


def phase1_payload(data: dict[str, Any]) -> dict[str, Any]:
    """Return the user-approved Stage A fields in a stable shape."""
    courses = []
    for course in data.get("courses", []):
        poster_selection = (
            course.get("poster_template_selection")
            or course.get("template_selection")
        )
        courses.append({
            "course_key": course.get("course_key") or course_key_for(course),
            "class_name": course.get("class_name"),
            "schedule_title": course.get("schedule_title"),
            "course_title": course.get("course_title"),
            "teacher": course.get("teacher"),
            "delivery": course.get("delivery"),
            "time": course.get("time"),
            "schedule_source": course.get("schedule_source"),
            "schedule_cell_locator": course.get("schedule_cell_locator"),
            "schedule_entry_index": course.get("schedule_entry_index"),
            "schedule_entry_count": course.get("schedule_entry_count"),
            "schedule_cell_color": course.get("schedule_cell_color"),
            "schedule_legend_evidence": course.get("schedule_legend_evidence"),
            "delivery_evidence_source": course.get("delivery_evidence_source"),
            "delivery_evidence": course.get("delivery_evidence"),
            "requested_poster_template_id": course.get("requested_poster_template_id"),
            "poster_template_selection": poster_selection,
        })
    return {
        "schema_version": data.get("schema_version"),
        "target_date": data.get("target_date"),
        "period": data.get("period"),
        "requested_text_template_id": data.get("requested_text_template_id"),
        "text_template_selection": data.get("text_template_selection"),
        "courses": courses,
    }


def phase1_fingerprint(data: dict[str, Any]) -> str:
    raw = json.dumps(
        phase1_payload(data),
        ensure_ascii=False,
        sort_keys=True,
        separators=(",", ":"),
    ).encode("utf-8")
    return hashlib.sha256(raw).hexdigest()


def file_sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def template_assets(args: argparse.Namespace) -> dict[str, Any]:
    return {
        "document_template": args.document_template,
        "poster_sources": list(args.poster_source or []),
    }


def check_template_id(template_id: str) -> None:
    if not TEMPLATE_ID_RE.fullmatch(template_id):
        raise ValueError("template id must use lowercase letters, digits, and single hyphens")


def read_registry(path: Path) -> dict[str, Any]:
    registry = read_json(path)
    if registry.get("schema_version") != 1:
        raise ValueError("template registry schema_version must be 1")
    if not isinstance(registry.get("templates"), dict) or not registry["templates"]:
        raise ValueError("template registry has no templates")
    default_id = registry.get("default_template_id")
    if default_id not in registry["templates"]:
        raise ValueError("default_template_id is missing from templates")
    return registry


def registry_template(registry: dict[str, Any], template_id: str) -> dict[str, Any]:
    template = registry["templates"].get(template_id)
    if not template:
        raise ValueError(f"unknown template id: {template_id}")
    if not template.get("enabled", True):
        raise ValueError(f"template is disabled: {template_id}")
    return template


def validate_asset_paths(registry: dict[str, Any], assets: dict[str, Any]) -> None:
    root = Path(registry["project_root"])
    values = []
    if assets.get("document_template"):
        values.append(assets["document_template"])
    values.extend(assets.get("poster_sources") or [])
    missing = []
    for value in values:
        path = Path(value)
        resolved = path if path.is_absolute() else root / path
        if not resolved.exists():
            missing.append(str(resolved))
    if missing:
        raise ValueError("template assets do not exist: " + " | ".join(missing))


def save_registry(path: Path, registry: dict[str, Any], increment: bool = True) -> None:
    if increment:
        registry["revision"] = int(registry.get("revision", 0)) + 1
    write_json(path, registry)


def template_init(args: argparse.Namespace) -> None:
    path = Path(args.registry)
    if path.exists():
        raise ValueError(f"registry already exists: {path}")
    check_template_id(args.id)
    registry = {
        "schema_version": 1,
        "revision": 1,
        "project_root": str(Path(args.root).resolve()),
        "default_template_id": args.id,
        "templates": {
            args.id: {
                "display_name": args.name,
                "enabled": True,
                "assets": template_assets(args),
                "notes": args.notes or "",
            }
        },
        "rules": [],
    }
    validate_asset_paths(registry, registry["templates"][args.id]["assets"])
    save_registry(path, registry, increment=False)
    print(path.resolve())


def template_add(args: argparse.Namespace) -> None:
    path = Path(args.registry)
    registry = read_registry(path)
    check_template_id(args.id)
    if args.id in registry["templates"]:
        raise ValueError(f"template already exists: {args.id}")
    assets = template_assets(args)
    validate_asset_paths(registry, assets)
    registry["templates"][args.id] = {
        "display_name": args.name,
        "enabled": True,
        "assets": assets,
        "notes": args.notes or "",
    }
    save_registry(path, registry)
    print(json.dumps({
        "added": args.id,
        "default_template_id": registry["default_template_id"],
        "revision": registry["revision"],
    }, ensure_ascii=False, indent=2))


def template_set_default(args: argparse.Namespace) -> None:
    path = Path(args.registry)
    registry = read_registry(path)
    registry_template(registry, args.id)
    registry["default_template_id"] = args.id
    save_registry(path, registry)
    print(json.dumps({
        "default_template_id": args.id,
        "display_name": registry["templates"][args.id]["display_name"],
        "revision": registry["revision"],
    }, ensure_ascii=False, indent=2))


def rule_match(rule: dict[str, Any], course: dict[str, Any]) -> bool:
    match = rule["match"]
    if match.get("course_key") and match["course_key"] != normalize_title(course.get("course_title", "")):
        return False
    if match.get("class_key") and match["class_key"] != normalize_title(course.get("class_name", "")):
        return False
    if match.get("delivery") and match["delivery"] != course.get("delivery"):
        return False
    return True


def rule_specificity(rule: dict[str, Any]) -> int:
    match = rule["match"]
    return (
        (4 if match.get("course_key") else 0)
        + (2 if match.get("class_key") else 0)
        + (1 if match.get("delivery") else 0)
    )


def resolve_template(
    registry: dict[str, Any],
    course: dict[str, Any],
    explicit_template_id: str | None = None,
) -> dict[str, Any]:
    if explicit_template_id:
        template = registry_template(registry, explicit_template_id)
        return {
            "template_id": explicit_template_id,
            "display_name": template["display_name"],
            "selection_source": "explicit",
            "registry_revision": registry["revision"],
        }
    matches = [rule for rule in registry.get("rules", []) if rule_match(rule, course)]
    if matches:
        best = max(rule_specificity(rule) for rule in matches)
        winners = [rule for rule in matches if rule_specificity(rule) == best]
        template_ids = {rule["template_id"] for rule in winners}
        if len(template_ids) > 1:
            raise ValueError(
                f"equal-specificity template conflict for {course.get('course_title')}: "
                + ", ".join(sorted(template_ids))
            )
        template_id = next(iter(template_ids))
        template = registry_template(registry, template_id)
        return {
            "template_id": template_id,
            "display_name": template["display_name"],
            "selection_source": "rule:" + ",".join(rule["rule_id"] for rule in winners),
            "registry_revision": registry["revision"],
        }
    template_id = registry["default_template_id"]
    template = registry_template(registry, template_id)
    return {
        "template_id": template_id,
        "display_name": template["display_name"],
        "selection_source": "default",
        "registry_revision": registry["revision"],
    }


def template_assign(args: argparse.Namespace) -> None:
    path = Path(args.registry)
    registry = read_registry(path)
    registry_template(registry, args.id)
    if args.delivery and args.delivery not in DELIVERIES:
        raise ValueError("delivery must be live or recorded")
    if not any((args.course_title, args.class_name, args.delivery)):
        raise ValueError("at least one rule selector is required")
    selector_match = {
        "course_title": args.course_title,
        "course_key": normalize_title(args.course_title or ""),
        "class_name": args.class_name,
        "class_key": normalize_title(args.class_name or ""),
        "delivery": args.delivery,
    }
    selector = {
        "course_key": selector_match["course_key"],
        "class_key": selector_match["class_key"],
        "delivery": selector_match["delivery"],
    }
    existing = [
        rule for rule in registry.get("rules", [])
        if {key: rule["match"].get(key) for key in selector} == selector
    ]
    if existing and not args.replace:
        raise ValueError("an identical selector already exists; use --replace to change it")
    if existing:
        registry["rules"] = [rule for rule in registry["rules"] if rule not in existing]
    used_numbers = [
        int(rule_match_obj.group(1))
        for rule in registry.get("rules", [])
        if (rule_match_obj := re.fullmatch(r"rule-(\d+)", rule.get("rule_id", "")))
    ]
    rule = {
        "rule_id": f"rule-{max(used_numbers, default=0) + 1:04d}",
        "template_id": args.id,
        "match": selector_match,
    }
    registry.setdefault("rules", []).append(rule)
    save_registry(path, registry)
    print(json.dumps(rule, ensure_ascii=False, indent=2))


def template_resolve(args: argparse.Namespace) -> None:
    registry = read_registry(Path(args.registry))
    course = {
        "course_title": args.course_title,
        "class_name": args.class_name or "",
        "delivery": args.delivery,
    }
    resolved = resolve_template(registry, course, args.explicit_template)
    resolved["assets"] = registry["templates"][resolved["template_id"]]["assets"]
    print(json.dumps(resolved, ensure_ascii=False, indent=2))


def template_list(args: argparse.Namespace) -> None:
    registry = read_registry(Path(args.registry))
    print(json.dumps(registry, ensure_ascii=False, indent=2))


def template_plan(args: argparse.Namespace) -> None:
    """Resolve the legacy combined registry while emitting schema-v2 fields.

    New runs should use workflow_guardrails.py template-plan with the separate
    text and poster registries. This compatibility command prevents an older
    manifest from becoming unbuildable during migration.
    """
    registry = read_registry(Path(args.registry))
    manifest = read_json(Path(args.manifest))
    previous = {
        "text": manifest.get("text_template_selection"),
        "posters": [
            course.get("poster_template_selection")
            or course.get("template_selection")
            for course in manifest.get("courses", [])
        ],
    }
    for course in manifest.get("courses", []):
        resolved = resolve_template(
            registry,
            course,
            course.get("requested_poster_template_id")
            or course.get("requested_template_id"),
        )
        course["poster_template_selection"] = resolved
        course["template_selection"] = resolved
    requested_document = (
        manifest.get("requested_text_template_id")
        or manifest.get("requested_document_template_id")
    )
    document_id = requested_document or registry["default_template_id"]
    document_template = registry_template(registry, document_id)
    manifest["document_template_id"] = document_id
    manifest["text_template_selection"] = {
        "template_id": document_id,
        "display_name": document_template["display_name"],
        "selection_source": "explicit" if requested_document else "default",
        "registry_revision": registry["revision"],
        "legacy_combined_registry": True,
    }
    current = {
        "text": manifest.get("text_template_selection"),
        "posters": [
            course.get("poster_template_selection")
            for course in manifest.get("courses", [])
        ],
    }
    if previous != current:
        approval = manifest.setdefault("approval", {})
        approval["phase1_status"] = "pending"
        approval["approved_at"] = None
        approval["phase1_content_hash"] = None
        approval["poster_status"] = "pending"
        approval["poster_approved_at"] = None
        approval["poster_content_hash"] = None
    write_json(Path(args.out), manifest)
    print(Path(args.out).resolve())


def init_manifest(args: argparse.Namespace) -> None:
    target = parse_date(args.date)
    root = Path(args.root).resolve()
    docs = Path(args.docs).resolve() if args.docs else (root / "课程预告文本").resolve()
    run_dir = (
        Path(args.run_dir).resolve()
        if args.run_dir
        else (root / "_course_preview_runs" / target.isoformat()).resolve()
    )
    period = period_for(target)
    if docs.is_dir():
        existing = existing_weekly_for(docs, target)
        if existing:
            period = existing
    source_document = period.get("existing_document")
    canonical_path = docs / period["filename"]
    weekly_document = {
        "canonical_path": str(canonical_path),
        "mode": "update" if source_document else "create",
        "source_document": source_document,
        "source_is_legacy_name": bool(
            source_document and Path(source_document).resolve() != canonical_path.resolve()
        ),
        "working_path": str(run_dir / "drafts" / "weekly-working.docx"),
        "backup_root": str(run_dir / "backups" / "weekly-document"),
        "same_week_candidates": (
            [
                {
                    "path": str(Path(source_document).resolve()),
                    "name": Path(source_document).name,
                    "size": Path(source_document).stat().st_size,
                    "sha256": file_sha256(Path(source_document)),
                }
            ]
            if source_document
            else []
        ),
        "publication_status": "pending",
        "published_sha256": None,
    }
    data = {
        "schema_version": SCHEMA_VERSION,
        "target_date": target.isoformat(),
        "root": str(root),
        "period": period,
        "weekly_document": weekly_document,
        "approval": {
            "phase1_status": "pending",
            "approved_at": None,
            "phase1_content_hash": None,
            "poster_status": "pending",
            "poster_approved_at": None,
            "poster_content_hash": None,
        },
        "requested_text_template_id": None,
        "text_template_selection": None,
        "requested_document_template_id": None,
        "document_template_id": None,
        "courses": [],
    }
    write_json(Path(args.out), data)
    print(Path(args.out).resolve())


def validate_manifest(data: dict[str, Any], phase: int = 1) -> list[str]:
    errors: list[str] = []
    if data.get("schema_version") != SCHEMA_VERSION:
        errors.append(f"schema_version must be {SCHEMA_VERSION}")
    try:
        parse_date(str(data["target_date"]))
    except Exception as exc:
        errors.append(f"invalid target_date: {exc}")
    courses = data.get("courses")
    if not isinstance(courses, list) or not courses:
        errors.append("courses must be a non-empty list")
        return errors
    if phase >= 2:
        approval = data.get("approval", {})
        if approval.get("phase1_status") != "approved":
            errors.append("approval.phase1_status must be approved for phase 2")
        approved_hash = approval.get("phase1_content_hash")
        if not approved_hash:
            errors.append("approval.phase1_content_hash is required for phase 2")
        elif approved_hash != phase1_fingerprint(data):
            errors.append(
                "Stage A content changed after approval; obtain approval again"
            )
    seen: set[tuple[str, str]] = set()
    seen_course_keys: set[str] = set()
    schedule_groups: dict[str, list[tuple[int, int, int]]] = {}
    for i, course in enumerate(courses):
        prefix = f"courses[{i}]"
        for field in ("class_name", "schedule_title", "course_title", "teacher",
                      "delivery", "schedule_source", "schedule_evidence", "confidence"):
            if not str(course.get(field, "")).strip():
                errors.append(f"{prefix}.{field} is required")
        if course.get("delivery") not in DELIVERIES:
            errors.append(f"{prefix}.delivery must be live or recorded")
        entry_index = course.get("schedule_entry_index")
        entry_count = course.get("schedule_entry_count")
        if not isinstance(entry_index, int) or isinstance(entry_index, bool) or entry_index < 1:
            errors.append(f"{prefix}.schedule_entry_index must be a positive integer")
        if not isinstance(entry_count, int) or isinstance(entry_count, bool) or entry_count < 1:
            errors.append(f"{prefix}.schedule_entry_count must be a positive integer")
        if (
            isinstance(entry_index, int)
            and not isinstance(entry_index, bool)
            and isinstance(entry_count, int)
            and not isinstance(entry_count, bool)
            and entry_index > entry_count
        ):
            errors.append(f"{prefix}.schedule_entry_index exceeds schedule_entry_count")
        for field in ("schedule_cell_locator", "schedule_cell_color", "delivery_evidence"):
            if not str(course.get(field, "")).strip():
                errors.append(f"{prefix}.{field} is required")
        delivery_source = course.get("delivery_evidence_source")
        if delivery_source not in DELIVERY_EVIDENCE_SOURCES:
            errors.append(
                f"{prefix}.delivery_evidence_source must be one of "
                f"{sorted(DELIVERY_EVIDENCE_SOURCES)}"
            )
        if delivery_source == "schedule_legend" and not str(
            course.get("schedule_legend_evidence", "")
        ).strip():
            errors.append(f"{prefix}.schedule_legend_evidence is required for schedule_legend")
        if delivery_source == "user_confirmation" and course.get("confidence") != "user_confirmed":
            errors.append(
                f"{prefix}.confidence must be user_confirmed when delivery came from the user"
            )
        group = str(course.get("class_name") or "").strip()
        if isinstance(entry_index, int) and isinstance(entry_count, int):
            schedule_groups.setdefault(group, []).append((i, entry_index, entry_count))
        expected_course_key = course_key_for(course)
        supplied_course_key = str(course.get("course_key") or "").strip()
        if supplied_course_key and supplied_course_key != expected_course_key:
            errors.append(
                f"{prefix}.course_key does not match normalized class/title"
            )
        course["course_key"] = expected_course_key
        if expected_course_key in seen_course_keys:
            errors.append(f"{prefix}.course_key is duplicated: {expected_course_key}")
        seen_course_keys.add(expected_course_key)
        for override in course.get("overrides") or []:
            if override.get("field") == "delivery":
                if not override.get("time_confirmed"):
                    errors.append(
                        f"{prefix}.overrides delivery change requires time_confirmed"
                    )
                for field in (
                    "old_delivery",
                    "new_delivery",
                    "old_time",
                    "new_time",
                    "evidence",
                ):
                    if override.get(field) in (None, ""):
                        errors.append(
                            f"{prefix}.overrides delivery change requires {field}"
                        )
        selection = (
            course.get("poster_template_selection")
            or course.get("template_selection")
        )
        if not isinstance(selection, dict):
            errors.append(
                f"{prefix}.poster_template_selection is required; "
                "run the separated template plan"
            )
        else:
            for field in ("template_id", "display_name", "selection_source", "registry_revision"):
                if selection.get(field) in (None, ""):
                    errors.append(
                        f"{prefix}.poster_template_selection.{field} is required"
                    )
        key = (str(course.get("class_name")), normalize_title(str(course.get("course_title", ""))))
        if key in seen:
            errors.append(f"{prefix} duplicates class/title {key}")
        seen.add(key)
        if phase >= 2:
            if not str(course.get("time", "")).strip():
                errors.append(f"{prefix}.time is required for phase 2")
            details = course.get("details")
            if not isinstance(details, dict):
                errors.append(f"{prefix}.details is required for phase 2")
            else:
                for field in ("tools", "homework", "objectives", "source_document"):
                    if field not in details:
                        errors.append(f"{prefix}.details.{field} is required")
                if not details.get("objectives"):
                    errors.append(f"{prefix}.details.objectives must not be empty")
                sources = details.get("sources")
                if not isinstance(sources, dict):
                    errors.append(f"{prefix}.details.sources is required")
                else:
                    for field in ("tools", "homework", "objectives"):
                        evidence = sources.get(field)
                        if not isinstance(evidence, dict):
                            errors.append(
                                f"{prefix}.details.sources.{field} is required"
                            )
                        elif not evidence.get("source_type") or not evidence.get("path"):
                            errors.append(
                                f"{prefix}.details.sources.{field} requires source_type and path"
                            )
    for group, items in schedule_groups.items():
        counts = {item[2] for item in items}
        label = f"class={group!r}"
        if len(counts) != 1:
            errors.append(
                f"schedule extraction count disagrees within {label}: {sorted(counts)}"
            )
            continue
        expected_count = next(iter(counts))
        indices = [item[1] for item in items]
        expected_indices = list(range(1, expected_count + 1))
        if sorted(indices) != expected_indices:
            errors.append(
                f"schedule extraction is incomplete for {label}: "
                f"found indices {sorted(indices)}, expected {expected_indices}"
            )
    text_selection = data.get("text_template_selection")
    if not isinstance(text_selection, dict):
        errors.append(
            "text_template_selection is required; run the separated template plan"
        )
    else:
        for field in ("template_id", "display_name", "selection_source", "registry_revision"):
            if text_selection.get(field) in (None, ""):
                errors.append(f"text_template_selection.{field} is required")
    if not str(data.get("document_template_id") or "").strip():
        errors.append("document_template_id compatibility field is required")
    return errors


def validate_command(args: argparse.Namespace) -> None:
    data = read_json(Path(args.manifest))
    errors = validate_manifest(data, args.phase)
    if errors:
        print("\n".join(f"ERROR: {item}" for item in errors), file=sys.stderr)
        raise SystemExit(2)
    print("OK")


def clear_document_body(doc: Document) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def set_east_asia(run, name: str = "宋体") -> None:
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    fonts.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", name)


def add_class_header(doc: Document, class_name: str) -> None:
    p = doc.add_paragraph()
    match = re.search(r"（.+?）", class_name)
    parts = [(class_name, False)] if not match else [
        (class_name[:match.start()], False),
        (match.group(), True),
        (class_name[match.end():], False),
    ]
    for text, bold in parts:
        if not text:
            continue
        run = p.add_run(text)
        set_east_asia(run)
        run.font.size = Pt(14)
        run.bold = bold


def course_line(course: dict[str, Any], target: date) -> str:
    suffix = "   （直播）" if course["delivery"] == LIVE else ""
    return f"{target.month}.{target.day}   {course['course_title']}   {str(course['teacher']).strip()}{suffix}"


def add_course_line(doc: Document, course: dict[str, Any], target: date):
    p = doc.add_paragraph()
    run = p.add_run(course_line(course, target))
    set_east_asia(run)
    run.font.size = Pt(10.5)
    if course["delivery"] == LIVE:
        run.bold = True
        run.font.color.rgb = RGBColor(255, 0, 0)
    return p


def add_compact_spacer(doc: Document, size_pt: float = 2.0):
    """Add a visible but compact gap without consuming a full body-text line."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(size_pt)
    run = p.add_run(" ")
    run.font.size = Pt(size_pt)
    return p


def load_template(path: Path, mode: str) -> Document:
    if not path.exists():
        raise FileNotFoundError(path)
    doc = Document(path)
    if mode == "new":
        clear_document_body(doc)
    return doc


def group_courses_by_class(
    courses: list[dict[str, Any]],
) -> list[tuple[str, list[dict[str, Any]]]]:
    """Group courses by exact class name while preserving first-seen order."""
    groups: list[tuple[str, list[dict[str, Any]]]] = []
    positions: dict[str, int] = {}
    for course in courses:
        class_name = str(course["class_name"])
        if class_name not in positions:
            positions[class_name] = len(groups)
            groups.append((class_name, []))
        groups[positions[class_name]][1].append(course)
    return groups


def build_phase1(args: argparse.Namespace) -> None:
    manifest = read_json(Path(args.manifest))
    errors = validate_manifest(manifest, 1)
    if errors:
        raise ValueError("; ".join(errors))
    target = parse_date(manifest["target_date"])
    doc = load_template(Path(args.template), args.mode)
    if args.mode == "append":
        date_re = re.compile(rf"^\s*0?{target.month}\.0?{target.day}(?:\D|$)")
        duplicates = []
        for course in manifest["courses"]:
            title_key = normalize_title(course["course_title"])
            if any(
                date_re.search(paragraph.text)
                and title_key in normalize_title(paragraph.text)
                for paragraph in doc.paragraphs
            ):
                duplicates.append(course_line(course, target))
        if duplicates:
            raise ValueError("refusing duplicate rows: " + " | ".join(duplicates))
    for class_name, class_courses in group_courses_by_class(manifest["courses"]):
        add_class_header(doc, class_name)
        for course in class_courses:
            add_course_line(doc, course, target)
        doc.add_paragraph()
    out = Path(args.out)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    print(out.resolve())


def add_detail_block(
    doc: Document,
    course: dict[str, Any],
    target: date,
    include_class_header: bool = True,
) -> None:
    details = course["details"]
    if include_class_header:
        add_class_header(doc, course["class_name"])
    course_paragraph = add_course_line(doc, course, target)
    notification_paragraph = doc.add_paragraph(
        details.get("notification") or
        ("🙌直播预告通知" if course["delivery"] == LIVE else "🙌录播课程已更新")
    )
    # Keep only the short course-line/notification pair together. Applying
    # keep-with-next to an entire detail block creates an overlong pagination
    # chain; Word may then position its first lines above the page crop box.
    course_paragraph.paragraph_format.keep_with_next = True
    notification_paragraph.paragraph_format.keep_together = True
    add_compact_spacer(doc)
    time = str(course["time"]).replace(":", "：")
    if course["delivery"] == LIVE:
        doc.add_paragraph(f"👉{target.month}月{target.day}日  {time}上课")
        add_compact_spacer(doc)
        doc.add_paragraph("👉主题：")
    else:
        doc.add_paragraph(f"👉{target.month}月{target.day}号  {time}")
    add_compact_spacer(doc)
    title_paragraph = doc.add_paragraph(f"【{course['course_title']}】")
    title_spacer = add_compact_spacer(doc)
    title_paragraph.paragraph_format.keep_with_next = True
    title_spacer.paragraph_format.keep_with_next = True
    doc.add_paragraph(f"授课老师：{course['teacher']}")
    add_compact_spacer(doc)
    tools = "、".join(details.get("tools") or [])
    doc.add_paragraph(f"所需工具：{tools}" if tools else "所需工具：待补充")
    homework = details.get("homework") or []
    if homework:
        add_compact_spacer(doc)
        doc.add_paragraph("作业布置：")
        for item in homework:
            doc.add_paragraph(str(item))
    add_compact_spacer(doc)
    doc.add_paragraph("课程目标：")
    for item in details.get("objectives") or []:
        doc.add_paragraph(str(item))
    add_compact_spacer(doc)


def normalize_existing_pagination(doc: Document) -> None:
    """Remove legacy overlong keep chains while retaining explicit page breaks."""
    paragraphs = list(doc.paragraphs)
    for index, paragraph in enumerate(paragraphs):
        # None removes the XML property entirely. Writing False leaves a
        # <w:keepNext w:val="0"/> marker; some Word/PDF export paths still
        # paginate legacy documents unreliably when hundreds of those markers
        # remain in the file.
        paragraph.paragraph_format.keep_with_next = None
        paragraph.paragraph_format.keep_together = None
        # Convert blank manual-break paragraphs into page-start buffers.
        # Word may place the paragraph immediately after a hard break above the
        # PDF crop box; a 2 pt invisible buffer absorbs that page-start offset.
        if (
            not paragraph.text.strip()
            and 'w:type="page"' in paragraph._p.xml
        ):
            paragraph.clear()
            run = paragraph.add_run(" ")
            run.font.size = Pt(14)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = Pt(14)
            paragraph.paragraph_format.page_break_before = True
    # Restore only short, semantically safe pairs so headings and labels are
    # not orphaned at page bottoms without recreating a whole-block chain.
    remaining = list(doc.paragraphs)
    for index, paragraph in enumerate(remaining[:-1]):
        text = paragraph.text.strip()
        next_text = remaining[index + 1].text.strip()
        if re.match(r"^\d+\s*期.*班", text):
            # Do not use keep-with-next on class headings. Word can shift a
            # naturally paginated heading above the PDF crop box when that
            # property crosses a page boundary; orphan headings are handled by
            # the render guard and a targeted page-start buffer instead.
            paragraph.paragraph_format.keep_with_next = None
        elif (
            re.match(r"^\d{1,2}\.\d{1,2}\s+", text)
            and ("课程已更新" in next_text or "预告通知" in next_text)
        ):
            paragraph.paragraph_format.keep_with_next = True
        elif text.startswith("【") and next_text.startswith("授课老师："):
            paragraph.paragraph_format.keep_with_next = True
        elif text.startswith("【"):
            # Compact spacer paragraphs may sit between a title and teacher.
            lookahead = index + 1
            while lookahead < len(remaining) and not remaining[lookahead].text.strip():
                lookahead += 1
            if (
                lookahead < len(remaining)
                and remaining[lookahead].text.strip().startswith("授课老师：")
            ):
                for pair_index in range(index, lookahead):
                    remaining[pair_index].paragraph_format.keep_with_next = True


def build_phase2(args: argparse.Namespace) -> None:
    manifest = read_json(Path(args.manifest))
    errors = validate_manifest(manifest, 2)
    if errors:
        raise ValueError("; ".join(errors))
    doc = Document(Path(args.phase1))
    normalize_existing_pagination(doc)
    target = parse_date(manifest["target_date"])
    for group_index, (class_name, class_courses) in enumerate(
        group_courses_by_class(manifest["courses"])
    ):
        if group_index == 0:
            page_start_buffer = add_compact_spacer(doc, size_pt=14)
            page_start_buffer.paragraph_format.page_break_before = True
        add_class_header(doc, class_name)
        for course in class_courses:
            add_detail_block(
                doc,
                course,
                target,
                include_class_header=False,
            )
    out = Path(args.out)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    print(out.resolve())


def paragraph_texts(doc_path: Path) -> list[str]:
    return [p.text.strip() for p in Document(doc_path).paragraphs]


LIVE_MARKER_RE = re.compile(r"(?:（直播）|\(直播\))")


def historical_delivery_from_explicit_marker(
    texts: list[str], title_index: int
) -> tuple[str | None, str | None]:
    """Read explicit DOCX markers only; title words such as 直播间 are not evidence."""
    for line in reversed(texts[max(0, title_index - 8):title_index]):
        normalized = re.sub(r"\s+", "", line)
        if "录播课程已更新" in normalized:
            return RECORDED, "录播课程已更新"
        if "直播预告通知" in normalized:
            return LIVE, "直播预告通知"
        if LIVE_MARKER_RE.search(line):
            return LIVE, "（直播）"
    return None, None


def catalog(args: argparse.Namespace) -> None:
    docs = Path(args.docs)
    entries: list[dict[str, Any]] = []
    for path in sorted(docs.glob("*.docx")):
        if WEEKLY_RE.match(path.name) and not CANONICAL_WEEKLY_RE.match(path.name):
            print(
                f"WARN: skipped non-canonical weekly DOCX {path.name}",
                file=sys.stderr,
            )
            continue
        try:
            texts = paragraph_texts(path)
        except Exception as exc:
            print(f"WARN: skipped {path.name}: {exc}", file=sys.stderr)
            continue
        for i, text in enumerate(texts):
            match = BRACKET_RE.fullmatch(text)
            if not match:
                continue
            title = match.group(1).strip()
            lo, hi = max(0, i - 6), min(len(texts), i + 28)
            block = texts[lo:hi]
            teacher = None
            for line in block:
                teacher_match = TEACHER_RE.search(line)
                if teacher_match:
                    teacher = teacher_match.group(1).strip()
                    break
            delivery, delivery_marker = historical_delivery_from_explicit_marker(texts, i)
            entries.append({
                "title": title,
                "normalized_title": normalize_title(title),
                "teacher": teacher,
                "delivery": delivery,
                "delivery_evidence_source": (
                    "historical_docx_explicit_marker" if delivery else "unknown"
                ),
                "delivery_evidence": delivery_marker,
                "source_document": path.name,
                "paragraph_index": i,
                "context": block,
            })
    write_json(Path(args.out), {"schema_version": 1, "entries": entries})
    print(f"{Path(args.out).resolve()} ({len(entries)} entries)")


def match_catalog(args: argparse.Namespace) -> None:
    catalog_data = read_json(Path(args.catalog))
    query = normalize_title(args.title)
    grouped: dict[tuple[str, str, str], dict[str, Any]] = {}
    for entry in catalog_data.get("entries", []):
        score = SequenceMatcher(None, query, entry["normalized_title"]).ratio()
        if score >= args.threshold:
            key = (
                entry["normalized_title"],
                str(entry.get("teacher") or ""),
                str(entry.get("delivery") or ""),
            )
            item = grouped.setdefault(key, {
                "title": entry.get("title"),
                "teacher": entry.get("teacher"),
                "delivery": entry.get("delivery"),
                "score": round(score, 4),
                "sources": [],
            })
            item["score"] = max(item["score"], round(score, 4))
            item["sources"].append({
                "source_document": entry.get("source_document"),
                "paragraph_index": entry.get("paragraph_index"),
            })
    results = list(grouped.values())
    results.sort(key=lambda item: (
        -item["score"],
        str(item.get("teacher") or ""),
        str(item.get("delivery") or ""),
    ))
    print(json.dumps(results[:args.limit], ensure_ascii=False, indent=2))


def template_poster_files(registry: dict[str, Any], template_id: str) -> list[Path]:
    template = registry_template(registry, template_id)
    root = Path(registry["project_root"])
    files: list[Path] = []
    for value in template.get("assets", {}).get("poster_sources") or []:
        path = Path(value)
        resolved = path if path.is_absolute() else root / path
        if resolved.is_file() and resolved.suffix.lower() in {".psd", ".png"}:
            files.append(resolved)
        elif resolved.is_dir():
            files.extend(
                item for item in resolved.rglob("*")
                if item.is_file() and item.suffix.lower() in {".psd", ".png"}
            )
    return sorted(set(files))


def poster_plan(args: argparse.Namespace) -> None:
    manifest = read_json(Path(args.manifest))
    registry = read_registry(Path(args.registry))
    errors = validate_manifest(manifest, 2)
    unresolved = [
        f"{course.get('class_name')}: teacher/confidence unresolved"
        for course in manifest.get("courses", [])
        if course.get("confidence") not in {"confirmed", "user_confirmed"}
        or course.get("teacher") == "待确认"
    ]
    if errors or unresolved:
        raise ValueError("; ".join(errors + unresolved))
    plan = []
    for course in manifest.get("courses", []):
        selection = (
            course.get("poster_template_selection")
            or course.get("template_selection")
        )
        if not selection:
            raise ValueError(
                f"missing poster template selection for {course['course_title']}"
            )
        if selection["registry_revision"] != registry["revision"]:
            raise ValueError(
                f"template registry changed after planning {course['course_title']}; "
                "rerun template-plan and obtain approval again"
            )
        template_id = selection["template_id"]
        files = template_poster_files(registry, template_id)
        if not files:
            raise ValueError(f"template {template_id} has no poster source assets")
        query = normalize_title(course["course_title"])
        class_query = normalize_title(course["class_name"])
        teacher_query = normalize_title(course["teacher"])
        candidates = []
        for path in files:
            name_norm = normalize_title(path.stem)
            exact = bool(query and query in name_norm)
            class_match = bool(class_query and class_query in name_norm)
            teacher_match = bool(teacher_query and teacher_query in name_norm)
            score = 1.0 if exact else SequenceMatcher(None, query, name_norm).ratio()
            if score >= 0.55:
                candidates.append({
                    "path": str(path),
                    "type": path.suffix.lower().lstrip("."),
                    "score": round(score, 4),
                    "exact_title_in_filename": exact,
                    "class_match": class_match,
                    "teacher_match": teacher_match,
                })
        candidates.sort(key=lambda item: (
            not item["exact_title_in_filename"],
            not item["class_match"],
            not item["teacher_match"],
            item["type"] != "psd",
            -item["score"],
            item["path"],
        ))
        plan.append({
            "class_name": course["class_name"],
            "course_title": course["course_title"],
            "template_selection": selection,
            "candidates": candidates[:8],
        })
    print(json.dumps(plan, ensure_ascii=False, indent=2))


def week_range(args: argparse.Namespace) -> None:
    print(json.dumps(period_for(parse_date(args.date)), ensure_ascii=False, indent=2))


def parser() -> argparse.ArgumentParser:
    result = argparse.ArgumentParser(description=__doc__)
    sub = result.add_subparsers(dest="command", required=True)

    cmd = sub.add_parser("week-range")
    cmd.add_argument("--date", required=True)
    cmd.set_defaults(func=week_range)

    cmd = sub.add_parser("init-manifest")
    cmd.add_argument("--date", required=True)
    cmd.add_argument("--root", required=True)
    cmd.add_argument("--docs")
    cmd.add_argument("--run-dir")
    cmd.add_argument("--out", required=True)
    cmd.set_defaults(func=init_manifest)

    cmd = sub.add_parser("template-init")
    cmd.add_argument("--registry", required=True)
    cmd.add_argument("--root", required=True)
    cmd.add_argument("--id", required=True)
    cmd.add_argument("--name", required=True)
    cmd.add_argument("--document-template")
    cmd.add_argument("--poster-source", action="append")
    cmd.add_argument("--notes")
    cmd.set_defaults(func=template_init)

    cmd = sub.add_parser("template-add")
    cmd.add_argument("--registry", required=True)
    cmd.add_argument("--id", required=True)
    cmd.add_argument("--name", required=True)
    cmd.add_argument("--document-template")
    cmd.add_argument("--poster-source", action="append")
    cmd.add_argument("--notes")
    cmd.set_defaults(func=template_add)

    cmd = sub.add_parser("template-set-default")
    cmd.add_argument("--registry", required=True)
    cmd.add_argument("--id", required=True)
    cmd.set_defaults(func=template_set_default)

    cmd = sub.add_parser("template-assign")
    cmd.add_argument("--registry", required=True)
    cmd.add_argument("--id", required=True)
    cmd.add_argument("--course-title")
    cmd.add_argument("--class-name")
    cmd.add_argument("--delivery", choices=(LIVE, RECORDED))
    cmd.add_argument("--replace", action="store_true")
    cmd.set_defaults(func=template_assign)

    cmd = sub.add_parser("template-resolve")
    cmd.add_argument("--registry", required=True)
    cmd.add_argument("--course-title", required=True)
    cmd.add_argument("--class-name")
    cmd.add_argument("--delivery", choices=(LIVE, RECORDED))
    cmd.add_argument("--explicit-template")
    cmd.set_defaults(func=template_resolve)

    cmd = sub.add_parser("template-list")
    cmd.add_argument("--registry", required=True)
    cmd.set_defaults(func=template_list)

    cmd = sub.add_parser("template-plan")
    cmd.add_argument("--registry", required=True)
    cmd.add_argument("--manifest", required=True)
    cmd.add_argument("--out", required=True)
    cmd.set_defaults(func=template_plan)

    cmd = sub.add_parser("validate")
    cmd.add_argument("--manifest", required=True)
    cmd.add_argument("--phase", type=int, choices=(1, 2), default=1)
    cmd.set_defaults(func=validate_command)

    cmd = sub.add_parser("build-phase1")
    cmd.add_argument("--manifest", required=True)
    cmd.add_argument("--template", required=True)
    cmd.add_argument("--mode", required=True, choices=("new", "append"))
    cmd.add_argument("--out", required=True)
    cmd.set_defaults(func=build_phase1)

    cmd = sub.add_parser("build-phase2")
    cmd.add_argument("--manifest", required=True)
    cmd.add_argument("--phase1", required=True)
    cmd.add_argument("--out", required=True)
    cmd.set_defaults(func=build_phase2)

    cmd = sub.add_parser("catalog")
    cmd.add_argument("--docs", required=True)
    cmd.add_argument("--out", required=True)
    cmd.set_defaults(func=catalog)

    cmd = sub.add_parser("match")
    cmd.add_argument("--catalog", required=True)
    cmd.add_argument("--title", required=True)
    cmd.add_argument("--threshold", type=float, default=0.72)
    cmd.add_argument("--limit", type=int, default=10)
    cmd.set_defaults(func=match_catalog)

    cmd = sub.add_parser("poster-plan")
    cmd.add_argument("--manifest", required=True)
    cmd.add_argument("--root", required=True)
    cmd.add_argument("--registry", required=True)
    cmd.set_defaults(func=poster_plan)
    return result


def main() -> None:
    args = parser().parse_args()
    try:
        args.func(args)
    except Exception as exc:
        print(f"ERROR: {exc}", file=sys.stderr)
        raise SystemExit(1)


if __name__ == "__main__":
    main()
