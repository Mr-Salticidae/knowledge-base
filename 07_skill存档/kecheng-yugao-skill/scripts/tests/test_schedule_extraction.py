# -*- coding: utf-8 -*-
from __future__ import annotations

import importlib.util
import tempfile
import unittest
from pathlib import Path

from docx import Document


SCRIPT = Path(__file__).resolve().parents[1] / "course_preview.py"
SPEC = importlib.util.spec_from_file_location("course_preview_schedule_test", SCRIPT)
assert SPEC and SPEC.loader
course_preview = importlib.util.module_from_spec(SPEC)
SPEC.loader.exec_module(course_preview)


def selection() -> dict:
    return {
        "template_id": "poster-template-1",
        "display_name": "海报模板一",
        "selection_source": "default",
        "registry_revision": 1,
    }


def course(title: str, index: int, count: int, delivery: str = "recorded") -> dict:
    return {
        "class_name": "22期 AIGC 凌云班（无基础）",
        "schedule_title": title,
        "course_title": title,
        "teacher": "禅基老师",
        "delivery": delivery,
        "schedule_source": "课表/设计班/凌云（无基础）.png",
        "schedule_evidence": f"8月3日第{index}个课程色块",
        "schedule_cell_locator": "八月第五周，周一3日",
        "schedule_entry_index": index,
        "schedule_entry_count": count,
        "schedule_cell_color": "青绿色",
        "schedule_legend_evidence": "绿色是直播，其他为录播",
        "delivery_evidence_source": "schedule_legend",
        "delivery_evidence": "青绿色属于图例中的其他颜色，因此为录播",
        "confidence": "confirmed",
        "poster_template_selection": selection(),
    }


def manifest(courses: list[dict]) -> dict:
    return {
        "schema_version": 2,
        "target_date": "2026-08-03",
        "courses": courses,
        "text_template_selection": {
            "template_id": "text-template-1",
            "display_name": "文本模板一",
            "selection_source": "default",
            "registry_revision": 1,
        },
        "document_template_id": "text-template-1",
    }


class ScheduleExtractionTests(unittest.TestCase):
    def test_accepts_all_courses_in_one_class_cell(self) -> None:
        data = manifest([
            course("全景概览与核心原理", 1, 2),
            course("可灵/即梦核心功能全解析", 2, 2),
        ])
        self.assertEqual(course_preview.validate_manifest(data, 1), [])

    def test_rejects_missing_second_course_in_class_cell(self) -> None:
        data = manifest([course("全景概览与核心原理", 1, 2)])
        errors = course_preview.validate_manifest(data, 1)
        self.assertTrue(any("schedule extraction is incomplete" in item for item in errors))

    def test_rejects_delivery_without_schedule_or_user_evidence(self) -> None:
        item = course("课程A", 1, 1)
        item["delivery_evidence_source"] = "historical_docx"
        errors = course_preview.validate_manifest(manifest([item]), 1)
        self.assertTrue(any("delivery_evidence_source" in value for value in errors))

    def test_title_containing_live_word_does_not_become_live(self) -> None:
        texts = [
            "21期 AIGC 灵机班（无基础）",
            "8.4 留人率飙升！Seedream打造高级直播间 Jelly老师",
            "🙌录播课程已更新",
            "【留人率飙升！Seedream打造高级直播间】",
        ]
        self.assertEqual(
            course_preview.historical_delivery_from_explicit_marker(texts, 3),
            ("recorded", "录播课程已更新"),
        )

    def test_no_explicit_docx_marker_stays_unknown(self) -> None:
        texts = [
            "直播贴片场景设计 Jelly老师",
            "【直播贴片场景设计】",
        ]
        self.assertEqual(
            course_preview.historical_delivery_from_explicit_marker(texts, 1),
            (None, None),
        )

    def test_catalog_keeps_unknown_instead_of_defaulting_recorded(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            doc = Document()
            doc.add_paragraph("直播贴片场景设计 Jelly老师")
            doc.add_paragraph("【直播贴片场景设计】")
            doc.save(root / "资料.docx")
            out = root / "catalog.json"
            args = type("Args", (), {"docs": str(root), "out": str(out)})()
            course_preview.catalog(args)
            data = course_preview.read_json(out)
            self.assertIsNone(data["entries"][0]["delivery"])
            self.assertEqual(data["entries"][0]["delivery_evidence_source"], "unknown")


if __name__ == "__main__":
    unittest.main()
