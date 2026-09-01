#!/usr/bin/env python3
"""Insert safe page-start buffers before render-identified orphan headings."""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.shared import Pt


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", required=True)
    parser.add_argument("--output", required=True)
    parser.add_argument("--indices", required=True, nargs="+", type=int)
    parser.add_argument("--buffer-pt", type=float, default=14.0)
    args = parser.parse_args()

    doc = Document(args.input)
    paragraphs = list(doc.paragraphs)
    for index in sorted(set(args.indices)):
        if index <= 0 or index >= len(paragraphs):
            raise ValueError(f"paragraph index out of range: {index}")
        heading = paragraphs[index]
        buffer_paragraph = paragraphs[index - 1]
        if buffer_paragraph.text.strip():
            raise ValueError(
                f"paragraph {index - 1} is not a blank buffer before {heading.text!r}"
            )
        buffer_paragraph.paragraph_format.space_before = Pt(0)
        buffer_paragraph.paragraph_format.space_after = Pt(0)
        buffer_paragraph.paragraph_format.line_spacing = Pt(args.buffer_pt)
        buffer_paragraph.paragraph_format.page_break_before = True
        buffer_paragraph.paragraph_format.keep_with_next = None
        if not buffer_paragraph.runs:
            buffer_paragraph.add_run(" ")
        for run in buffer_paragraph.runs:
            run.font.size = Pt(args.buffer_pt)
        heading.paragraph_format.keep_with_next = None

    output = Path(args.output)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    print(output.resolve())


if __name__ == "__main__":
    main()
